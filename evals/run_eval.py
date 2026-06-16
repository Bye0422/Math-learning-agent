import os
import sys
import csv
import json
import time
import uuid
import base64
import mimetypes
import re
from pathlib import Path

from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document as DocxReader

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma


# =========================
# 路径设置
# =========================

PROJECT_ROOT = Path(__file__).resolve().parents[1]
EVALS_DIR = PROJECT_ROOT / "evals"
SOURCE_DOCS_DIR = EVALS_DIR / "source_docs"
EVAL_CASES_PATH = EVALS_DIR / "eval_cases.json"
RESULTS_CSV_PATH = EVALS_DIR / "eval_results.csv"
SUMMARY_JSON_PATH = EVALS_DIR / "eval_summary.json"

sys.path.append(str(PROJECT_ROOT))


# =========================
# 导入项目已有模块
# =========================

from prompts.ocr_prompts import OCR_IMAGE_PROMPT
from prompts.router_prompts import build_task_classification_prompt
from prompts.retrieval_prompts import build_rewrite_query_prompt
from prompts.answer_prompts import build_rag_answer_prompt
from prompts.direct_prompts import build_direct_answer_prompt
from prompts.validation_prompts import build_answer_repair_prompt

from validators.task_validator import validate_task_info
from validators.answer_validator import validate_answer


# =========================
# 读取环境变量
# =========================

load_dotenv(PROJECT_ROOT / ".env")

QWEN_API_KEY = os.getenv("QWEN_API_KEY")
QWEN_BASE_URL = os.getenv("QWEN_BASE_URL")
QWEN_CHAT_MODEL = os.getenv("QWEN_CHAT_MODEL", "qwen3.7-plus")
QWEN_OCR_MODEL = os.getenv("QWEN_OCR_MODEL", QWEN_CHAT_MODEL)
QWEN_EMBEDDING_MODEL = os.getenv("QWEN_EMBEDDING_MODEL", "text-embedding-v4")


# =========================
# 基础检查
# =========================

def check_env():
    if not QWEN_API_KEY:
        raise RuntimeError("没有读取到 QWEN_API_KEY，请检查 .env 文件。")

    if not QWEN_BASE_URL:
        raise RuntimeError("没有读取到 QWEN_BASE_URL，请检查 .env 文件。")


def get_chat_llm(temperature=0.2):
    return ChatOpenAI(
        model=QWEN_CHAT_MODEL,
        api_key=QWEN_API_KEY,
        base_url=QWEN_BASE_URL,
        temperature=temperature,
    )


# =========================
# 文档读取
# =========================

def read_pdf_to_documents(file_path):
    documents = []

    with open(file_path, "rb") as f:
        reader = PdfReader(f)

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                print(f"[WARN] PDF 加密无法读取：{file_path.name}")
                return []

        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()

            if page_text and page_text.strip():
                documents.append(
                    Document(
                        page_content=page_text.strip(),
                        metadata={
                            "source": file_path.name,
                            "file_type": "pdf",
                            "location": f"第 {page_num} 页"
                        }
                    )
                )

    return documents


def read_docx_to_documents(file_path):
    docx = DocxReader(str(file_path))
    documents = []

    for para_num, paragraph in enumerate(docx.paragraphs, start=1):
        text = paragraph.text.strip()

        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "file_type": "docx",
                        "location": f"第 {para_num} 段"
                    }
                )
            )

    table_index = 0

    for table in docx.tables:
        table_index += 1
        rows_text = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

            if cells:
                rows_text.append(" | ".join(cells))

        if rows_text:
            documents.append(
                Document(
                    page_content="\n".join(rows_text),
                    metadata={
                        "source": file_path.name,
                        "file_type": "docx",
                        "location": f"表格 {table_index}"
                    }
                )
            )

    return documents


def read_txt_to_documents(file_path):
    raw_data = file_path.read_bytes()

    try:
        text = raw_data.decode("utf-8")
    except UnicodeDecodeError:
        text = raw_data.decode("gbk", errors="ignore")

    text = text.strip()

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "file_type": "txt",
                "location": "全文"
            }
        )
    ]


def image_file_to_data_url(file_path):
    image_bytes = file_path.read_bytes()

    mime_type, _ = mimetypes.guess_type(file_path.name)

    if mime_type is None:
        suffix = file_path.suffix.lower()

        if suffix == ".png":
            mime_type = "image/png"
        elif suffix in [".jpg", ".jpeg"]:
            mime_type = "image/jpeg"
        elif suffix == ".webp":
            mime_type = "image/webp"
        else:
            mime_type = "image/png"

    base64_image = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{base64_image}"


def ocr_image_with_qwen(file_path):
    data_url = image_file_to_data_url(file_path)

    llm = ChatOpenAI(
        model=QWEN_OCR_MODEL,
        api_key=QWEN_API_KEY,
        base_url=QWEN_BASE_URL,
        temperature=0.1,
    )

    message = HumanMessage(
        content=[
            {
                "type": "text",
                "text": OCR_IMAGE_PROMPT
            },
            {
                "type": "image_url",
                "image_url": {
                    "url": data_url
                }
            }
        ]
    )

    response = llm.invoke([message])
    return response.content.strip()


def read_image_to_documents(file_path):
    try:
        text = ocr_image_with_qwen(file_path)
    except Exception as e:
        print(f"[WARN] 图片 OCR 失败：{file_path.name}，原因：{e}")
        return []

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "file_type": "image",
                "location": "OCR识别文本"
            }
        )
    ]


def read_file_to_documents(file_path):
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            return read_pdf_to_documents(file_path)

        if suffix == ".docx":
            return read_docx_to_documents(file_path)

        if suffix == ".txt":
            return read_txt_to_documents(file_path)

        if suffix in [".png", ".jpg", ".jpeg", ".webp"]:
            return read_image_to_documents(file_path)

        print(f"[WARN] 暂不支持文件类型：{file_path.name}")
        return []

    except Exception as e:
        print(f"[WARN] 读取失败：{file_path.name}，原因：{e}")
        return []


def load_source_documents():
    if not SOURCE_DOCS_DIR.exists():
        raise FileNotFoundError(f"请先创建文件夹：{SOURCE_DOCS_DIR}")

    file_paths = [
        p for p in SOURCE_DOCS_DIR.iterdir()
        if p.is_file() and p.suffix.lower() in [".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".webp"]
    ]

    if not file_paths:
        raise FileNotFoundError(f"请先把评估文档放到：{SOURCE_DOCS_DIR}")

    all_documents = []

    for file_path in file_paths:
        print(f"[INFO] 正在读取：{file_path.name}")
        documents = read_file_to_documents(file_path)
        all_documents.extend(documents)

    if not all_documents:
        raise RuntimeError("没有读取到任何有效文档内容。")

    return all_documents


# =========================
# RAG 基础流程
# =========================

def split_documents(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=120,
        length_function=len,
        separators=["\n\n", "\n", "。", "，", " ", ""]
    )

    chunks = text_splitter.split_documents(documents)

    clean_chunks = []

    for idx, chunk in enumerate(chunks, start=1):
        if chunk.page_content and chunk.page_content.strip():
            chunk.page_content = chunk.page_content.strip()
            chunk.metadata["chunk_id"] = idx
            clean_chunks.append(chunk)

    return clean_chunks


def create_vector_db(chunks):
    embeddings = OpenAIEmbeddings(
        model=QWEN_EMBEDDING_MODEL,
        api_key=QWEN_API_KEY,
        base_url=QWEN_BASE_URL,
        tiktoken_enabled=False,
        check_embedding_ctx_length=False,
        chunk_size=10,
    )

    collection_name = f"eval_rag_{uuid.uuid4().hex[:8]}"

    vector_db = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=collection_name
    )

    return vector_db


def search_similar_chunks_with_score(vector_db, question, k=3):
    return vector_db.similarity_search_with_score(question, k=k)


# =========================
# 对话、JSON、任务识别
# =========================

def format_chat_history(chat_history, max_messages=6):
    recent_history = chat_history[-max_messages:]
    history_text = ""

    for msg in recent_history:
        if msg["role"] == "user":
            history_text += f"用户：{msg['content']}\n"
        elif msg["role"] == "assistant":
            history_text += f"AI：{msg['content']}\n"

    return history_text.strip()


def parse_json_from_text(text):
    text = text.strip()

    text = re.sub(r"^```json", "", text)
    text = re.sub(r"^```", "", text)
    text = re.sub(r"```$", "", text)
    text = text.strip()

    try:
        return json.loads(text)
    except Exception:
        pass

    match = re.search(r"\{.*\}", text, re.DOTALL)

    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            pass

    return None


def classify_task(question, chat_history):
    history_text = format_chat_history(chat_history, max_messages=6)

    prompt = build_task_classification_prompt(
        question=question,
        history_text=history_text
    )

    try:
        llm = get_chat_llm(temperature=0.1)
        response = llm.invoke(prompt)
        raw_result = parse_json_from_text(response.content)
        return validate_task_info(raw_result)

    except Exception:
        return validate_task_info(None)


def rewrite_question_for_retrieval(question, chat_history, task_info):
    if not chat_history and task_info.get("task_type") != "summary":
        return question

    history_text = format_chat_history(chat_history, max_messages=6)
    task_type = task_info.get("task_type", "qa")

    prompt = build_rewrite_query_prompt(
        question=question,
        history_text=history_text,
        task_type=task_type
    )

    try:
        llm = get_chat_llm(temperature=0.1)
        response = llm.invoke(prompt)
        rewritten_question = response.content.strip()

        if rewritten_question:
            return rewritten_question

        return question

    except Exception:
        return question


# =========================
# 生成答案 + Guardrails
# =========================

def build_source_summary(docs):
    lines = []

    for i, doc in enumerate(docs, start=1):
        source = doc.metadata.get("source", "未知文件")
        file_type = doc.metadata.get("file_type", "未知类型")
        location = doc.metadata.get("location", "未知位置")
        chunk_id = doc.metadata.get("chunk_id", "未知片段")

        lines.append(
            f"资料片段 {i}：{source}，{file_type}，{location}，Chunk {chunk_id}"
        )

    return "\n".join(lines)


def generate_answer_with_sources(question, docs, chat_history, task_info):
    history_text = format_chat_history(chat_history, max_messages=6)

    prompt = build_rag_answer_prompt(
        question=question,
        docs=docs,
        history_text=history_text,
        task_info=task_info
    )

    llm = get_chat_llm(temperature=0.2)
    response = llm.invoke(prompt)
    answer = response.content

    task_type = task_info.get("task_type", "qa")
    validation_result = validate_answer(answer, task_type)

    was_repaired = False

    if validation_result["is_valid"]:
        return answer, validation_result, was_repaired

    source_summary = build_source_summary(docs)

    repair_prompt = build_answer_repair_prompt(
        question=question,
        original_answer=answer,
        task_info=task_info,
        validation_result=validation_result,
        source_summary=source_summary
    )

    repair_llm = get_chat_llm(temperature=0.1)
    repair_response = repair_llm.invoke(repair_prompt)
    repaired_answer = repair_response.content

    repaired_validation_result = validate_answer(repaired_answer, task_type)
    was_repaired = True

    return repaired_answer, repaired_validation_result, was_repaired


def generate_direct_answer(question, chat_history, task_info):
    history_text = format_chat_history(chat_history, max_messages=6)

    prompt = build_direct_answer_prompt(
        question=question,
        history_text=history_text,
        task_info=task_info
    )

    llm = get_chat_llm(temperature=0.2)
    response = llm.invoke(prompt)

    return response.content


# =========================
# Eval 工具函数
# =========================

def load_eval_cases():
    if not EVAL_CASES_PATH.exists():
        raise FileNotFoundError(f"找不到评估集文件：{EVAL_CASES_PATH}")

    with open(EVAL_CASES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def keyword_score(text, keywords):
    if not keywords:
        return 1.0, 0, 0

    text = text or ""
    hit_count = 0

    for keyword in keywords:
        if keyword in text:
            hit_count += 1

    total = len(keywords)
    score = hit_count / total if total else 1.0

    return score, hit_count, total


def bool_to_int(value):
    if value is None:
        return ""

    return 1 if value else 0


def run_single_case(case, vector_db):
    start_time = time.time()

    case_id = case.get("id", "")
    question = case.get("question", "")
    chat_history = case.get("chat_history", [])

    expected_task_type = case.get("expected_task_type")
    expected_need_rag = case.get("expected_need_rag")

    expected_answer_keywords = case.get("expected_answer_keywords", [])
    expected_source_keywords = case.get("expected_source_keywords", [])

    result_row = {
        "id": case_id,
        "question": question,
        "expected_task_type": expected_task_type,
        "actual_task_type": "",
        "task_type_correct": "",
        "expected_need_rag": expected_need_rag,
        "actual_need_rag": "",
        "need_rag_correct": "",
        "retrieval_question": "",
        "retrieval_hit": "",
        "retrieval_keyword_score": "",
        "retrieval_keyword_hit_count": "",
        "retrieval_keyword_total": "",
        "answer_hit": "",
        "answer_keyword_score": "",
        "answer_keyword_hit_count": "",
        "answer_keyword_total": "",
        "format_valid": "",
        "was_repaired": "",
        "elapsed_seconds": "",
        "answer": "",
        "error": ""
    }

    try:
        task_info = classify_task(question, chat_history)

        actual_task_type = task_info.get("task_type", "")
        actual_need_rag = task_info.get("need_rag", True)

        result_row["actual_task_type"] = actual_task_type
        result_row["actual_need_rag"] = actual_need_rag

        if expected_task_type is not None:
            result_row["task_type_correct"] = bool_to_int(actual_task_type == expected_task_type)

        if expected_need_rag is not None:
            result_row["need_rag_correct"] = bool_to_int(actual_need_rag == expected_need_rag)

        if not actual_need_rag:
            answer = generate_direct_answer(question, chat_history, task_info)

            answer_score, answer_hit_count, answer_total = keyword_score(
                answer,
                expected_answer_keywords
            )

            result_row["answer"] = answer
            result_row["answer_keyword_score"] = round(answer_score, 4)
            result_row["answer_keyword_hit_count"] = answer_hit_count
            result_row["answer_keyword_total"] = answer_total
            result_row["answer_hit"] = bool_to_int(answer_score == 1.0)

            result_row["format_valid"] = ""
            result_row["was_repaired"] = 0

        else:
            retrieval_question = rewrite_question_for_retrieval(
                question,
                chat_history,
                task_info
            )

            result_row["retrieval_question"] = retrieval_question

            if actual_task_type in ["summary", "extract_points"]:
                top_k = case.get("top_k", 5)
            else:
                top_k = case.get("top_k", 3)

            results_with_scores = search_similar_chunks_with_score(
                vector_db,
                retrieval_question,
                k=top_k
            )

            docs = [doc for doc, score in results_with_scores]

            retrieved_context = "\n\n".join([doc.page_content for doc in docs])

            retrieval_score, retrieval_hit_count, retrieval_total = keyword_score(
                retrieved_context,
                expected_source_keywords
            )

            if expected_source_keywords:
                result_row["retrieval_hit"] = bool_to_int(retrieval_score == 1.0)
                result_row["retrieval_keyword_score"] = round(retrieval_score, 4)
                result_row["retrieval_keyword_hit_count"] = retrieval_hit_count
                result_row["retrieval_keyword_total"] = retrieval_total
            else:
                result_row["retrieval_hit"] = ""
                result_row["retrieval_keyword_score"] = ""
                result_row["retrieval_keyword_hit_count"] = ""
                result_row["retrieval_keyword_total"] = ""

            answer, validation_result, was_repaired = generate_answer_with_sources(
                question,
                docs,
                chat_history,
                task_info
            )

            answer_score, answer_hit_count, answer_total = keyword_score(
                answer,
                expected_answer_keywords
            )

            result_row["answer"] = answer
            result_row["answer_keyword_score"] = round(answer_score, 4)
            result_row["answer_keyword_hit_count"] = answer_hit_count
            result_row["answer_keyword_total"] = answer_total
            result_row["answer_hit"] = bool_to_int(answer_score == 1.0)

            result_row["format_valid"] = bool_to_int(validation_result.get("is_valid", False))
            result_row["was_repaired"] = bool_to_int(was_repaired)

        elapsed_seconds = time.time() - start_time
        result_row["elapsed_seconds"] = round(elapsed_seconds, 2)

        return result_row

    except Exception as e:
        elapsed_seconds = time.time() - start_time
        result_row["elapsed_seconds"] = round(elapsed_seconds, 2)
        result_row["error"] = str(e)
        return result_row


def write_results_csv(rows):
    if not rows:
        return

    fieldnames = list(rows[0].keys())

    with open(RESULTS_CSV_PATH, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def safe_rate(rows, field):
    values = []

    for row in rows:
        value = row.get(field, "")

        if value == "":
            continue

        try:
            values.append(int(value))
        except Exception:
            continue

    if not values:
        return None

    return round(sum(values) / len(values), 4)


def safe_average(rows, field):
    values = []

    for row in rows:
        value = row.get(field, "")

        if value == "":
            continue

        try:
            values.append(float(value))
        except Exception:
            continue

    if not values:
        return None

    return round(sum(values) / len(values), 4)


def build_summary(rows):
    summary = {
        "total_cases": len(rows),
        "task_type_accuracy": safe_rate(rows, "task_type_correct"),
        "need_rag_accuracy": safe_rate(rows, "need_rag_correct"),
        "retrieval_hit_rate": safe_rate(rows, "retrieval_hit"),
        "answer_keyword_hit_rate": safe_rate(rows, "answer_hit"),
        "format_valid_rate": safe_rate(rows, "format_valid"),
        "repair_trigger_rate": safe_rate(rows, "was_repaired"),
        "avg_retrieval_keyword_score": safe_average(rows, "retrieval_keyword_score"),
        "avg_answer_keyword_score": safe_average(rows, "answer_keyword_score"),
        "avg_elapsed_seconds": safe_average(rows, "elapsed_seconds"),
        "error_count": sum(1 for row in rows if row.get("error"))
    }

    return summary


def write_summary_json(summary):
    with open(SUMMARY_JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)


def print_summary(summary):
    print("\n========== Eval Summary ==========")
    print(f"测试样本数：{summary['total_cases']}")
    print(f"任务识别准确率：{summary['task_type_accuracy']}")
    print(f"Need RAG 判断准确率：{summary['need_rag_accuracy']}")
    print(f"检索命中率：{summary['retrieval_hit_rate']}")
    print(f"答案关键词命中率：{summary['answer_keyword_hit_rate']}")
    print(f"格式合规率：{summary['format_valid_rate']}")
    print(f"自动修复触发率：{summary['repair_trigger_rate']}")
    print(f"平均检索关键词得分：{summary['avg_retrieval_keyword_score']}")
    print(f"平均答案关键词得分：{summary['avg_answer_keyword_score']}")
    print(f"平均耗时：{summary['avg_elapsed_seconds']} 秒")
    print(f"错误数：{summary['error_count']}")
    print("==================================\n")


# =========================
# 主流程
# =========================

def main():
    check_env()

    print("[INFO] 开始加载评估文档...")
    documents = load_source_documents()
    print(f"[INFO] 原始 Document 数量：{len(documents)}")

    print("[INFO] 开始切分 chunk...")
    chunks = split_documents(documents)
    print(f"[INFO] chunk 数量：{len(chunks)}")

    print("[INFO] 开始建立向量库...")
    vector_db = create_vector_db(chunks)
    print("[INFO] 向量库建立完成。")

    print("[INFO] 开始加载评估样本...")
    eval_cases = load_eval_cases()
    print(f"[INFO] 评估样本数量：{len(eval_cases)}")

    rows = []

    for index, case in enumerate(eval_cases, start=1):
        case_id = case.get("id", f"case_{index}")
        question = case.get("question", "")

        print(f"\n[INFO] 正在评估 {index}/{len(eval_cases)}：{case_id}")
        print(f"[QUESTION] {question}")

        row = run_single_case(case, vector_db)
        rows.append(row)

        print(f"[RESULT] task={row.get('actual_task_type')} answer_hit={row.get('answer_hit')} format_valid={row.get('format_valid')} error={row.get('error')}")

    write_results_csv(rows)
    summary = build_summary(rows)
    write_summary_json(summary)
    print_summary(summary)

    print(f"[INFO] 详细结果已保存：{RESULTS_CSV_PATH}")
    print(f"[INFO] 汇总结果已保存：{SUMMARY_JSON_PATH}")


if __name__ == "__main__":
    main()