import base64
import mimetypes

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage

from config import USE_MINERU_FOR_PDF

from prompts.ocr_prompts import OCR_IMAGE_PROMPT
from services.llm_service import get_ocr_llm
from services.mineru_loader import read_pdf_with_mineru


def read_pdf_with_pypdf(uploaded_file):
    """
    使用 pypdf 读取 PDF。
    这是 fallback 方案：当 MinerU 失败时使用。
    """
    documents = []
    errors = []

    try:
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                errors.append(f"{uploaded_file.name} 是加密文件，需要密码才能读取。")
                return documents, errors

        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()

            if page_text and page_text.strip():
                documents.append(
                    Document(
                        page_content=page_text.strip(),
                        metadata={
                            "source": uploaded_file.name,
                            "file_type": "pdf_pypdf",
                            "location": f"第 {page_num} 页",
                        },
                    )
                )

        return documents, errors

    except Exception as e:
        errors.append(f"{uploaded_file.name} 使用 pypdf 读取失败：{e}")
        return documents, errors


def read_pdf_to_documents(uploaded_file):
    """
    PDF 读取入口。

    优先使用 MinerU：
    - MinerU 成功：返回 MinerU Markdown Document
    - MinerU 失败：自动回退到 pypdf
    """
    all_errors = []

    if USE_MINERU_FOR_PDF:
        mineru_documents, mineru_errors = read_pdf_with_mineru(uploaded_file)
        all_errors.extend(mineru_errors)

        if mineru_documents:
            return mineru_documents, all_errors

    pypdf_documents, pypdf_errors = read_pdf_with_pypdf(uploaded_file)
    all_errors.extend(pypdf_errors)

    return pypdf_documents, all_errors


def read_docx_to_documents(uploaded_file):
    """
    读取 Word .docx 文件，并把段落和表格转成 Document。
    """
    documents = []
    errors = []

    try:
        uploaded_file.seek(0)
        docx = DocxDocument(uploaded_file)

        for para_num, paragraph in enumerate(docx.paragraphs, start=1):
            text = paragraph.text.strip()

            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": uploaded_file.name,
                            "file_type": "docx",
                            "location": f"第 {para_num} 段",
                        },
                    )
                )

        table_index = 0

        for table in docx.tables:
            table_index += 1
            rows_text = []

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

                if cells:
                    rows_text.append(" | ".join(cells))

            if rows_text:
                table_text = "\n".join(rows_text)

                documents.append(
                    Document(
                        page_content=table_text,
                        metadata={
                            "source": uploaded_file.name,
                            "file_type": "docx",
                            "location": f"表格 {table_index}",
                        },
                    )
                )

        return documents, errors

    except Exception as e:
        errors.append(f"{uploaded_file.name} 读取失败：{e}")
        return documents, errors


def read_txt_to_documents(uploaded_file):
    """
    读取 TXT 文件，优先 UTF-8，失败后尝试 GBK。
    """
    documents = []
    errors = []

    try:
        uploaded_file.seek(0)
        raw_data = uploaded_file.read()

        try:
            text = raw_data.decode("utf-8")
        except UnicodeDecodeError:
            text = raw_data.decode("gbk", errors="ignore")

        text = text.strip()

        if not text:
            return documents, errors

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": uploaded_file.name,
                    "file_type": "txt",
                    "location": "全文",
                },
            )
        )

        return documents, errors

    except Exception as e:
        errors.append(f"{uploaded_file.name} 读取失败：{e}")
        return documents, errors


def image_to_data_url(uploaded_file):
    """
    把上传图片转成 base64 data URL。
    """
    uploaded_file.seek(0)
    image_bytes = uploaded_file.read()

    mime_type, _ = mimetypes.guess_type(uploaded_file.name)

    if mime_type is None:
        file_name = uploaded_file.name.lower()

        if file_name.endswith(".png"):
            mime_type = "image/png"
        elif file_name.endswith(".jpg") or file_name.endswith(".jpeg"):
            mime_type = "image/jpeg"
        elif file_name.endswith(".webp"):
            mime_type = "image/webp"
        else:
            mime_type = "image/png"

    base64_image = base64.b64encode(image_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{base64_image}"


def ocr_image_with_qwen(uploaded_file):
    """
    使用千问视觉模型进行 OCR。
    """
    data_url = image_to_data_url(uploaded_file)
    llm = get_ocr_llm()

    message = HumanMessage(
        content=[
            {
                "type": "text",
                "text": OCR_IMAGE_PROMPT,
            },
            {
                "type": "image_url",
                "image_url": {
                    "url": data_url,
                },
            },
        ]
    )

    response = llm.invoke([message])
    return response.content.strip()


def read_image_to_documents(uploaded_file):
    """
    图片 OCR 后转成 Document。
    """
    documents = []
    errors = []

    try:
        text = ocr_image_with_qwen(uploaded_file)

        if not text:
            return documents, errors

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": uploaded_file.name,
                    "file_type": "image",
                    "location": "OCR识别文本",
                },
            )
        )

        return documents, errors

    except Exception as e:
        errors.append(
            f"{uploaded_file.name} OCR 失败：{e}。如果报模型不支持图片，请把 .env 里的 QWEN_OCR_MODEL 改成支持视觉识别的模型。"
        )
        return documents, errors


def read_file_to_documents(uploaded_file):
    """
    根据文件后缀自动路由到不同解析器。
    """
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return read_pdf_to_documents(uploaded_file)

    if file_name.endswith(".docx"):
        return read_docx_to_documents(uploaded_file)

    if file_name.endswith(".txt"):
        return read_txt_to_documents(uploaded_file)

    if (
        file_name.endswith(".png")
        or file_name.endswith(".jpg")
        or file_name.endswith(".jpeg")
        or file_name.endswith(".webp")
    ):
        return read_image_to_documents(uploaded_file)

    return [], [f"暂不支持该文件类型：{uploaded_file.name}"]


def read_multiple_files_to_documents(uploaded_files):
    """
    读取多个文件，并统一合并成 Document 列表。
    """
    all_documents = []
    all_errors = []

    for uploaded_file in uploaded_files:
        documents, errors = read_file_to_documents(uploaded_file)
        all_documents.extend(documents)
        all_errors.extend(errors)

    return all_documents, all_errors