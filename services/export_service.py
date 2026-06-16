from datetime import datetime
from io import BytesIO

from docx import Document as DocxDocument

from config import EXPORT_TITLE


def build_export_text(chat_history, file_names, session_id, last_log_row=None):
    """
    构造 TXT 导出文本。
    """
    lines = []

    lines.append(EXPORT_TITLE)
    lines.append("=" * 40)
    lines.append(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"Session ID：{session_id}")
    lines.append("")

    lines.append("上传文件：")
    if file_names:
        for file_name in file_names:
            lines.append(f"- {file_name}")
    else:
        lines.append("- 无")
    lines.append("")

    lines.append("问答记录：")
    lines.append("-" * 40)

    if not chat_history:
        lines.append("暂无问答记录。")
    else:
        round_index = 1

        for msg in chat_history:
            role = msg.get("role", "")
            content = msg.get("content", "")

            if role == "user":
                lines.append("")
                lines.append(f"【第 {round_index} 轮】")
                lines.append(f"用户：{content}")

            elif role == "assistant":
                lines.append("")
                lines.append(f"AI：{content}")
                lines.append("-" * 40)
                round_index += 1

    if last_log_row:
        lines.append("")
        lines.append("最近一次运行日志摘要：")
        lines.append("-" * 40)
        lines.append(f"任务类型：{last_log_row.get('task_type', '')}")
        lines.append(f"是否需要 RAG：{last_log_row.get('need_rag', '')}")
        lines.append(f"实际检索问题：{last_log_row.get('retrieval_question', '')}")
        lines.append(f"Top K：{last_log_row.get('top_k', '')}")
        lines.append(f"格式是否合格：{last_log_row.get('format_valid', '')}")
        lines.append(f"是否触发修复：{last_log_row.get('was_repaired', '')}")
        lines.append(f"耗时：{last_log_row.get('elapsed_seconds', '')} 秒")
        lines.append(f"错误信息：{last_log_row.get('error', '')}")

    return "\n".join(lines)


def build_export_docx_bytes(chat_history, file_names, session_id, last_log_row=None):
    """
    构造 Word 导出文件 bytes。
    """
    doc = DocxDocument()

    doc.add_heading(EXPORT_TITLE, level=1)

    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Session ID：{session_id}")

    doc.add_heading("上传文件", level=2)

    if file_names:
        for file_name in file_names:
            doc.add_paragraph(file_name, style="List Bullet")
    else:
        doc.add_paragraph("无")

    doc.add_heading("问答记录", level=2)

    if not chat_history:
        doc.add_paragraph("暂无问答记录。")
    else:
        round_index = 1

        for msg in chat_history:
            role = msg.get("role", "")
            content = msg.get("content", "")

            if role == "user":
                doc.add_heading(f"第 {round_index} 轮", level=3)
                p = doc.add_paragraph()
                p.add_run("用户：").bold = True
                p.add_run(content)

            elif role == "assistant":
                p = doc.add_paragraph()
                p.add_run("AI：").bold = True
                p.add_run(content)
                round_index += 1

    if last_log_row:
        doc.add_heading("最近一次运行日志摘要", level=2)

        log_items = [
            ("任务类型", last_log_row.get("task_type", "")),
            ("是否需要 RAG", last_log_row.get("need_rag", "")),
            ("实际检索问题", last_log_row.get("retrieval_question", "")),
            ("Top K", last_log_row.get("top_k", "")),
            ("格式是否合格", last_log_row.get("format_valid", "")),
            ("是否触发修复", last_log_row.get("was_repaired", "")),
            ("耗时", f"{last_log_row.get('elapsed_seconds', '')} 秒"),
            ("错误信息", last_log_row.get("error", "")),
        ]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "字段"
        header_cells[1].text = "内容"

        for key, value in log_items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()